from dataclasses import replace
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

from md2docx.config import (
    REQUIRED_SECTIONS,
    Length,
    apply_config_to_style,
    list_level_layout,
    load_config,
)
from md2docx.converter import convert_markdown, parse_frontmatter
from md2docx.math_render import render_latex

PROJECT_ROOT = Path(__file__).parents[1]
CONFIG_PATH = PROJECT_ROOT / "config.yaml"
CONFIG = CONFIG_PATH.read_text(encoding="utf-8")


def test_frontmatter_is_optional() -> None:
    metadata, body = parse_frontmatter("# 普通标题")
    assert metadata == {}
    assert body == "# 普通标题"

    metadata, body = parse_frontmatter("---\ntitle: 文档标题\n---\n# 一级标题\n")
    assert metadata["title"] == "文档标题"
    assert body.startswith("# 一级标题")

    with pytest.raises(ValueError, match="closing"):
        parse_frontmatter("---\ntitle: 文档标题\n")


def test_real_config_is_the_strict_schema(tmp_path: Path) -> None:
    styles = load_config(CONFIG_PATH)
    assert set(styles) == REQUIRED_SECTIONS
    assert styles["title"].first_line_indent is None
    assert styles["body"].first_line_indent is not None
    assert styles["body"].first_line_indent.unit == "em"
    assert styles["title"].space_before.unit == "pt"
    assert styles["title"].space_after.unit == "pt"
    assert styles["title"].indent_before_text.unit == "pt"
    assert styles["ordered-list"].indent_before_text_increment is not None
    assert styles["unordered-list"].indent_before_text_increment is not None
    assert styles["ordered-list"].hanging_indent is not None

    config = tmp_path / "config.yaml"
    config.write_text(
        CONFIG.replace("inline-code:", "missing-inline-code:", 1), encoding="utf-8"
    )
    with pytest.raises(
        ValueError, match="missing required configuration section.*inline-code"
    ):
        load_config(config)

    config.write_text(
        CONFIG.replace('  color: "#000000"\n', "", 1),
        encoding="utf-8",
    )
    with pytest.raises(ValueError, match="title is missing required field.*color"):
        load_config(config)

    config.write_text(
        CONFIG.replace(
            "  hanging-indent: null\n  first-line-indent: null",
            '  hanging-indent: "1em"\n  first-line-indent: "1em"',
            1,
        ),
        encoding="utf-8",
    )
    with pytest.raises(ValueError, match="mutually exclusive"):
        load_config(config)


def test_end_to_end_uses_yaml_named_styles_without_direct_formatting(
    tmp_path: Path,
) -> None:
    Image.new("RGB", (320, 160), "steelblue").save(tmp_path / "sample.png")
    markdown = tmp_path / "sample.md"
    markdown.write_text(
        """---
title: 测试文档
---

# 概述

正文含有 **粗体**、*斜体*、~~删除线~~、`code` 和行内公式 $E=mc^2$。

1. 有序一层
   1. 有序二层
      1. 有序三层

- 无序一层
  - 无序二层

```text
code block
```

$$
\\frac{a}{b} = c
$$

![示例图片](sample.png)

| 左对齐 | 居中 | 右对齐 |
| :--- | :---: | ---: |
| 文本 | 文本 | 文本 |
""",
        encoding="utf-8",
    )
    output = tmp_path / "sample.docx"
    convert_markdown(markdown, output, CONFIG_PATH)

    document = Document(output)
    assert document.paragraphs[0].style.name == "title"
    assert next(p for p in document.paragraphs if p.text == "概述").style.name == "h1"
    assert (
        next(p for p in document.paragraphs if p.text == "code block").style.name
        == "code-block"
    )

    ordered = [
        p for p in document.paragraphs if p.style.name.startswith("ordered-list-")
    ]
    unordered = [
        p for p in document.paragraphs if p.style.name.startswith("unordered-list-")
    ]
    assert [p.style.name for p in ordered] == [
        "ordered-list-1",
        "ordered-list-2",
        "ordered-list-3",
    ]
    assert [p.style.name for p in unordered] == [
        "unordered-list-1",
        "unordered-list-2",
    ]
    list_config = load_config(CONFIG_PATH)["ordered-list"]
    assert [
        round(document.styles[f"ordered-list-{level}"].paragraph_format.left_indent.pt)
        for level in range(1, 4)
    ] == [
        round(list_level_layout(list_config, level).left_indent_pt)
        for level in range(1, 4)
    ]
    assert [
        round(
            document.styles[f"unordered-list-{level}"].paragraph_format.left_indent.pt
        )
        for level in range(1, 3)
    ] == [
        round(list_level_layout(list_config, level).left_indent_pt)
        for level in range(1, 3)
    ]
    assert [
        round(
            document.styles[
                f"ordered-list-{level}"
            ].paragraph_format.first_line_indent.pt
        )
        for level in range(1, 4)
    ] == [-32, -32, -32]

    table = document.tables[0]
    assert [cell.paragraphs[0].style.name for cell in table.rows[0].cells] == [
        "table-header",
        "table-header",
        "table-header",
    ]
    assert [cell.paragraphs[0].style.name for cell in table.rows[1].cells] == [
        "table-body",
        "table-body",
        "table-body",
    ]
    for row in table.rows:
        assert [cell.paragraphs[0].alignment for cell in row.cells] == [
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ]

    body = next(p for p in document.paragraphs if p.text.startswith("正文含有"))
    deleted = next(run for run in body.runs if run.text == "删除线")
    assert deleted.font.strike is True
    for run in body.runs:
        rpr = run._r.rPr
        if rpr is None:
            continue
        tags = {child.tag.rsplit("}", 1)[-1] for child in rpr}
        assert tags <= {"b", "bCs", "i", "iCs", "strike", "rStyle", "drawing"}

    for paragraph in document.paragraphs:
        if paragraph.text and paragraph.style.name not in {"body"}:
            ppr = paragraph._p.pPr
            if ppr is not None:
                tags = {child.tag.rsplit("}", 1)[-1] for child in ppr}
                assert tags <= {"pStyle"}

    with ZipFile(output) as archive:
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")

    for name in REQUIRED_SECTIONS - {"ordered-list", "unordered-list"}:
        assert f'w:name w:val="{name}"' in styles_xml
    assert 'w:name w:val="ordered-list"' not in styles_xml
    assert 'w:name w:val="unordered-list"' not in styles_xml
    for level in range(1, 10):
        assert f'w:name w:val="ordered-list-{level}"' in styles_xml
        assert f'w:name w:val="unordered-list-{level}"' in styles_xml
    assert "md2docx ordered-list numbering" in numbering_xml
    assert "md2docx unordered-list numbering" in numbering_xml
    assert 'w:pStyle w:val="ordered-list-1"' in numbering_xml
    assert 'w:pStyle w:val="unordered-list-1"' in numbering_xml
    for level in range(1, 4):
        layout = list_level_layout(list_config, level)
        assert f'w:left="{round(layout.left_indent_pt * 20)}"' in numbering_xml
        assert (
            f'w:hanging="{round(layout.hanging_indent_pt * 20)}"'
            in numbering_xml
        )
    assert "<w:rFonts" not in document_xml
    assert "<w:sz " not in document_xml
    assert "<w:color " not in document_xml
    assert "<w:strike" in document_xml


def test_inline_renderer_preserves_formatting_inside_hyperlinks(
    tmp_path: Path,
) -> None:
    markdown = tmp_path / "link.md"
    markdown.write_text(
        "访问 [**粗体**、`代码`、~~删除~~](https://example.com)。\n",
        encoding="utf-8",
    )
    output = tmp_path / "link.docx"

    convert_markdown(markdown, output, CONFIG_PATH)

    document = Document(output)
    paragraph = document.paragraphs[0]
    assert paragraph.text == "访问 粗体、代码、删除。"
    hyperlink = paragraph._p.find(qn("w:hyperlink"))
    assert hyperlink is not None
    runs = hyperlink.findall(qn("w:r"))
    assert "".join(run.find(qn("w:t")).text for run in runs) == "粗体、代码、删除"
    for run in runs:
        properties = run.find(qn("w:rPr"))
        assert properties.find(qn("w:color")).get(qn("w:val")) == "0563C1"
        assert properties.find(qn("w:u")).get(qn("w:val")) == "single"
    assert runs[0].find(qn("w:rPr") + "/" + qn("w:b")) is not None
    assert runs[2].find(qn("w:rPr") + "/" + qn("w:rStyle")) is not None
    assert runs[4].find(qn("w:rPr") + "/" + qn("w:strike")) is not None

    with ZipFile(output) as archive:
        relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
    assert 'Target="https://example.com"' in relationships
    assert 'TargetMode="External"' in relationships


def test_thematic_break_uses_native_paragraph_border(tmp_path: Path) -> None:
    markdown = tmp_path / "rule.md"
    markdown.write_text("上文\n\n---\n\n下文\n", encoding="utf-8")
    output = tmp_path / "rule.docx"

    convert_markdown(markdown, output, CONFIG_PATH)

    document = Document(output)
    horizontal_rule = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p.find(qn("w:pPr") + "/" + qn("w:pBdr")) is not None
    )
    bottom_border = horizontal_rule._p.find(
        qn("w:pPr") + "/" + qn("w:pBdr") + "/" + qn("w:bottom")
    )
    assert bottom_border is not None
    assert bottom_border.get(qn("w:val")) == "single"


def test_list_level_layout_is_the_single_indent_calculation() -> None:
    config = replace(
        load_config(CONFIG_PATH)["ordered-list"],
        indent_before_text=Length(10, "pt"),
        hanging_indent=Length(5, "pt"),
        indent_before_text_increment=Length(3, "pt"),
    )

    assert [
        list_level_layout(config, level).left_indent_pt for level in range(1, 4)
    ] == [15, 18, 21]
    assert [
        list_level_layout(config, level).hanging_indent_pt
        for level in range(1, 4)
    ] == [5, 5, 5]
    with pytest.raises(ValueError, match="at least 1"):
        list_level_layout(config, 0)


def test_markdown_without_frontmatter_has_no_title_paragraph(tmp_path: Path) -> None:
    markdown = tmp_path / "no-frontmatter.md"
    markdown.write_text("# 正文标题\n\n正文。\n", encoding="utf-8")
    output = tmp_path / "no-frontmatter.docx"

    convert_markdown(markdown, output, CONFIG_PATH)

    document = Document(output)
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "正文标题",
        "正文。",
    ]
    assert all(paragraph.style.name != "title" for paragraph in document.paragraphs)


def test_space_before_and_after_support_em_lengths(tmp_path: Path) -> None:
    config = tmp_path / "config.yaml"
    config.write_text(
        CONFIG.replace('  space-before: "28pt"', '  space-before: "1em"', 1).replace(
            '  space-after: "28pt"', '  space-after: "0.5em"', 1
        ),
        encoding="utf-8",
    )

    styles = load_config(config)
    document = Document()

    apply_config_to_style(document.styles["Title"], styles["title"])
    assert document.styles["Title"].paragraph_format.space_before.pt == 22
    assert document.styles["Title"].paragraph_format.space_after.pt == 11


def test_local_image_path_supports_chinese_characters(tmp_path: Path) -> None:
    image_name = "结算包拓扑关系示意图.png"
    Image.new("RGB", (32, 16), "steelblue").save(tmp_path / image_name)
    markdown = tmp_path / "中文文档.md"
    markdown.write_text(
        f"---\ntitle: 中文图片路径测试\n---\n\n![拓扑关系]({image_name})\n",
        encoding="utf-8",
    )
    output = tmp_path / "输出文档.docx"
    convert_markdown(markdown, output, CONFIG_PATH)
    with ZipFile(output) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())


def test_svg_is_embedded_directly_without_raster_fallback(tmp_path: Path) -> None:
    svg = b"""<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 200 100">
<rect width="200" height="100" fill="steelblue"/>
</svg>"""
    (tmp_path / "diagram.svg").write_bytes(svg)
    markdown = tmp_path / "svg.md"
    markdown.write_text("![SVG diagram](diagram.svg)\n", encoding="utf-8")
    output = tmp_path / "svg.docx"

    convert_markdown(markdown, output, CONFIG_PATH)

    with ZipFile(output) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        assert media == ["word/media/image1.svg"]
        assert archive.read(media[0]) == svg
        content_types = archive.read("[Content_Types].xml").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")
        relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")

    assert 'ContentType="image/svg+xml"' in content_types
    assert "svgBlip" in document_xml
    assert "drawing/2016/SVG/main" in document_xml
    assert 'Target="media/image1.svg"' in relationships
    assert not any(name.endswith((".png", ".jpg", ".jpeg")) for name in media)
    Document(output)


def test_math_fontset_renders_without_configuring_font_family(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    configured_rc: dict[str, str] = {}
    original_rc_context = __import__("matplotlib").rc_context

    def capture_rc_context(rc: dict[str, str]):
        configured_rc.update(rc)
        return original_rc_context(rc)

    monkeypatch.setattr("md2docx.math_render.mpl.rc_context", capture_rc_context)
    image = render_latex("E=mc^2", fontset="cm")
    assert image.read(8) == b"\x89PNG\r\n\x1a\n"
    assert configured_rc == {"mathtext.fontset": "cm"}


def test_unknown_math_fontset_fails_instead_of_silently_falling_back() -> None:
    with pytest.raises(ValueError, match="unsupported math fontset"):
        render_latex("x", fontset="Definitely Missing Math Font")
