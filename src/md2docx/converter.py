from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
from typing import Any, Iterable
from urllib.parse import unquote, urlparse
from urllib.request import urlopen
from xml.etree import ElementTree

import mistune
import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shape import CT_Inline
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.shared import Inches, Pt
from lxml import etree
from PIL import Image

from .config import (
    INLINE_SECTIONS,
    LIST_SECTIONS,
    StyleConfig,
    apply_config_to_style,
    load_config,
)
from .math_render import render_latex
from .numbering import (
    install_caption_numbering,
    install_heading_numbering,
    install_list_numbering,
    set_style_numbering,
)


def parse_frontmatter(source: str) -> tuple[dict[str, Any], str]:
    lines = source.splitlines(keepends=True)
    if not lines or lines[0].strip() != "---":
        return {}, source
    end = next(
        (index for index, line in enumerate(lines[1:], start=1) if line.strip() == "---"),
        None,
    )
    if end is None:
        raise ValueError("Markdown Frontmatter is missing its closing '---'")
    try:
        metadata = yaml.safe_load("".join(lines[1:end]))
    except yaml.YAMLError as exc:
        raise ValueError(f"invalid Markdown Frontmatter YAML: {exc}") from exc
    if not isinstance(metadata, dict):
        raise ValueError("Markdown Frontmatter must be a mapping")
    title = metadata.get("title")
    if not isinstance(title, str) or not title.strip():
        raise ValueError("Markdown Frontmatter must contain a non-empty title")
    return metadata, "".join(lines[end + 1 :])


def _reset_style(style: Any) -> None:
    for tag in ("w:basedOn", "w:link", "w:pPr", "w:rPr"):
        child = style.element.find(qn(tag))
        if child is not None:
            style.element.remove(child)


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, value: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


_SVG_NAMESPACE = "http://www.w3.org/2000/svg"
_SVG_BLIP_NAMESPACE = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
_SVG_EXTENSION_URI = "{96DAC541-7B7A-43D3-8B79-37D633B846F1}"
_RELATIONSHIP_NAMESPACE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
)


def _svg_length_px(value: str | None) -> float | None:
    if value is None:
        return None
    match = re.fullmatch(
        r"\s*(\d+(?:\.\d+)?|\.\d+)\s*(px|pt|pc|in|cm|mm)?\s*",
        value,
        flags=re.IGNORECASE,
    )
    if match is None:
        return None
    number = float(match.group(1))
    unit = (match.group(2) or "px").lower()
    pixels_per_unit = {
        "px": 1.0,
        "pt": 96 / 72,
        "pc": 16.0,
        "in": 96.0,
        "cm": 96 / 2.54,
        "mm": 96 / 25.4,
    }
    return number * pixels_per_unit[unit]


def _svg_dimensions(svg: bytes) -> tuple[float, float]:
    try:
        root = ElementTree.fromstring(svg)
    except ElementTree.ParseError as exc:
        raise ValueError(f"invalid SVG image: {exc}") from exc
    if root.tag not in {"svg", f"{{{_SVG_NAMESPACE}}}svg"}:
        raise ValueError("invalid SVG image: root element must be <svg>")

    width = _svg_length_px(root.get("width"))
    height = _svg_length_px(root.get("height"))
    view_box = root.get("viewBox")
    view_width = view_height = None
    if view_box is not None:
        try:
            _, _, view_width, view_height = map(
                float, re.split(r"[\s,]+", view_box.strip())
            )
        except ValueError as exc:
            raise ValueError("invalid SVG viewBox") from exc
        if view_width <= 0 or view_height <= 0:
            raise ValueError("SVG viewBox dimensions must be positive")

    if width is None and height is None:
        width, height = view_width, view_height
    elif width is None and height is not None and view_width and view_height:
        width = height * view_width / view_height
    elif height is None and width is not None and view_width and view_height:
        height = width * view_height / view_width

    if width is None or height is None or width <= 0 or height <= 0:
        raise ValueError("SVG requires positive width/height or a valid viewBox")
    return width, height


class DocxBuilder:
    def __init__(self, styles: dict[str, StyleConfig], source_dir: Path):
        self.styles = styles
        self.source_dir = source_dir
        self.document = Document()
        self._configure_document()

    def style(self, name: str) -> StyleConfig:
        try:
            return self.styles[name]
        except KeyError as exc:
            raise ValueError(f"missing required configuration section: {name}") from exc

    def _ensure_style(self, name: str, style_type: WD_STYLE_TYPE) -> Any:
        if name in self.document.styles:
            word_style = self.document.styles[name]
            if word_style.type != style_type:
                raise ValueError(f"Word style {name} has the wrong style type")
        else:
            word_style = self.document.styles.add_style(name, style_type)
        _reset_style(word_style)
        apply_config_to_style(word_style, self.style(name))
        return word_style

    def _configure_document(self) -> None:
        section = self.document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for name in self.styles:
            if name in LIST_SECTIONS:
                continue
            style_type = (
                WD_STYLE_TYPE.CHARACTER
                if name in INLINE_SECTIONS
                else WD_STYLE_TYPE.PARAGRAPH
            )
            self._ensure_style(name, style_type)

        headings = sorted(
            (
                (int(name[1:]), config)
                for name, config in self.styles.items()
                if re.fullmatch(r"h[1-9]", name)
            ),
            key=lambda item: item[0],
        )
        heading_num_id = install_heading_numbering(self.document, headings)
        for level, config in headings:
            if config.numbering is not None:
                set_style_numbering(self.document.styles[config.name], heading_num_id, level)

        caption = self.style("image-caption")
        set_style_numbering(
            self.document.styles[caption.name],
            install_caption_numbering(self.document, caption),
            1,
        )

        for name, ordered in (("ordered-list", True), ("unordered-list", False)):
            config = self.style(name)
            num_id = install_list_numbering(self.document, config, ordered=ordered)
            if config.indent_before_text_increment is None:
                raise ValueError(f"{name} requires indent-before-text-increment")
            for level in range(1, 10):
                style_name = f"{name}-{level}"
                word_style = self.document.styles.add_style(
                    style_name, WD_STYLE_TYPE.PARAGRAPH
                )
                apply_config_to_style(word_style, config)
                hanging_indent = (
                    0
                    if config.hanging_indent is None
                    else config.hanging_indent.to_points(config.size_pt)
                )
                word_style.paragraph_format.left_indent = Pt(
                    config.indent_before_text.to_points(config.size_pt)
                    + hanging_indent
                    + config.indent_before_text_increment.to_points(config.size_pt)
                    * (level - 1)
                )
                set_style_numbering(word_style, num_id, level)

    def add_title(self, title: str) -> None:
        self.document.add_paragraph(title.strip(), style="title")

    def add_heading(self, token: dict[str, Any]) -> None:
        name = f"h{int(token['attrs']['level'])}"
        paragraph = self.document.add_paragraph(style=name)
        self.add_inline_nodes(paragraph, token.get("children", []))

    def add_paragraph(self, children: Iterable[dict[str, Any]]) -> None:
        children = list(children)
        if len(children) == 1 and children[0]["type"] == "image":
            self.add_figure(children[0])
            return
        paragraph = self.document.add_paragraph(style="body")
        self.add_inline_nodes(paragraph, children)

    def add_inline_nodes(
        self,
        paragraph: Any,
        nodes: Iterable[dict[str, Any]],
        *,
        bold: bool = False,
        italic: bool = False,
        strike: bool = False,
    ) -> None:
        for node in nodes:
            kind = node["type"]
            if kind == "text":
                run = paragraph.add_run(node.get("raw", ""))
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                if strike:
                    run.font.strike = True
            elif kind in {"strong", "emphasis"}:
                self.add_inline_nodes(
                    paragraph,
                    node.get("children", []),
                    bold=bold or kind == "strong",
                    italic=italic or kind == "emphasis",
                    strike=strike,
                )
            elif kind == "codespan":
                run = paragraph.add_run(node.get("raw", ""))
                run.style = "inline-code"
            elif kind == "inline_math":
                config = self.style("inline-math")
                image = render_latex(
                    node.get("raw", ""),
                    font_size=config.size_pt or 16,
                    fontset=config.latin_font or "stix",
                    color=f"#{config.color}" if config.color is not None else "black",
                )
                paragraph.add_run().add_picture(image)
            elif kind == "image":
                self._insert_image(paragraph.add_run(), node)
            elif kind == "link":
                paragraph.add_run(
                    self._plain_text(node.get("children", []))
                    or node.get("attrs", {}).get("url", "")
                )
            elif kind in {"linebreak", "softbreak"}:
                paragraph.add_run().add_break()
            elif kind == "strikethrough":
                self.add_inline_nodes(
                    paragraph,
                    node.get("children", []),
                    bold=bold,
                    italic=italic,
                    strike=True,
                )
            else:
                self.add_inline_nodes(
                    paragraph,
                    node.get("children", []),
                    bold=bold,
                    italic=italic,
                    strike=strike,
                )

    def add_block_math(self, expression: str) -> None:
        config = self.style("math-block")
        paragraph = self.document.add_paragraph(style="math-block")
        image = render_latex(
            expression,
            font_size=config.size_pt or 16,
            fontset=config.latin_font or "stix",
            color=f"#{config.color}" if config.color is not None else "black",
        )
        paragraph.add_run().add_picture(image)

    def _image_stream(self, url: str) -> BytesIO:
        parsed = urlparse(url)
        if parsed.scheme in {"http", "https"}:
            with urlopen(url, timeout=20) as response:
                return BytesIO(response.read())
        path = Path(unquote(url))
        if not path.is_absolute():
            path = self.source_dir / path
        if not path.is_file():
            raise FileNotFoundError(f"image not found: {path}")
        return BytesIO(path.read_bytes())

    def _insert_svg(
        self,
        run: Any,
        svg: bytes,
        filename: str,
        *,
        max_width: float,
    ) -> None:
        width_px, height_px = _svg_dimensions(svg)
        width_in = min(width_px / 96, max_width)
        height_in = width_in * height_px / width_px

        document_part = run.part
        package = document_part.package
        partname = package.next_partname("/word/media/image%d.svg")
        svg_part = Part(partname, "image/svg+xml", svg, package)
        relationship_id = document_part.relate_to(svg_part, RT.IMAGE)

        inline = CT_Inline.new_pic_inline(
            document_part.next_id,
            relationship_id,
            filename or partname.filename,
            Inches(width_in),
            Inches(height_in),
        )
        blip = inline.find(".//" + qn("a:blip"))
        if blip is None:
            raise RuntimeError("failed to create SVG drawing")
        extension_list = etree.SubElement(blip, qn("a:extLst"))
        extension = etree.SubElement(extension_list, qn("a:ext"))
        extension.set("uri", _SVG_EXTENSION_URI)
        svg_blip = etree.SubElement(
            extension,
            f"{{{_SVG_BLIP_NAMESPACE}}}svgBlip",
            nsmap={"asvg": _SVG_BLIP_NAMESPACE},
        )
        svg_blip.set(f"{{{_RELATIONSHIP_NAMESPACE}}}embed", relationship_id)
        run._r.add_drawing(inline)

    def _insert_image(self, run: Any, token: dict[str, Any], *, block: bool = False) -> None:
        url = token.get("attrs", {}).get("url", "")
        stream = self._image_stream(url)
        max_width = 6.5 if block else 2.0
        image_bytes = stream.getvalue()
        try:
            root = ElementTree.fromstring(image_bytes)
            is_svg = root.tag in {"svg", f"{{{_SVG_NAMESPACE}}}svg"}
        except ElementTree.ParseError:
            is_svg = False
        if is_svg:
            filename = Path(unquote(urlparse(url).path)).name or "image.svg"
            self._insert_svg(run, image_bytes, filename, max_width=max_width)
            return

        with Image.open(stream) as image:
            width_px, height_px = image.size
            dpi = image.info.get("dpi", (96, 96))[0] or 96
            width_in = min(width_px / dpi, max_width)
            height_in = width_in * height_px / max(width_px, 1)
        stream.seek(0)
        run.add_picture(stream, width=Inches(width_in), height=Inches(height_in))

    def add_figure(self, token: dict[str, Any]) -> None:
        paragraph = self.document.add_paragraph(style="image")
        self._insert_image(paragraph.add_run(), token, block=True)
        caption = self._plain_text(token.get("children", []))
        if caption:
            self.document.add_paragraph(caption, style="image-caption")

    def add_table(self, token: dict[str, Any]) -> None:
        rows: list[list[dict[str, Any]]] = []
        header = next(
            (child for child in token["children"] if child["type"] == "table_head"),
            None,
        )
        body = next(
            (child for child in token["children"] if child["type"] == "table_body"),
            None,
        )
        if header:
            rows.append(header.get("children", []))
        if body:
            rows.extend(row.get("children", []) for row in body.get("children", []))
        if not rows:
            return
        column_count = max(len(row) for row in rows)
        table = self.document.add_table(rows=len(rows), cols=column_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths = [9360 // column_count] * column_count
        widths[-1] += 9360 - sum(widths)
        alignments = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        for row_index, row in enumerate(rows):
            for column_index, cell_token in enumerate(row):
                cell = table.cell(row_index, column_index)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_cell_margins(cell)
                paragraph = cell.paragraphs[0]
                paragraph.style = "table-header" if row_index == 0 else "table-body"
                alignment = cell_token.get("attrs", {}).get("align")
                if alignment is not None:
                    paragraph.alignment = alignments[alignment]
                self.add_inline_nodes(paragraph, cell_token.get("children", []))
                if row_index == 0:
                    _shade_cell(cell, "E7E6E6")
                    for run in paragraph.runs:
                        run.bold = True
        _set_table_geometry(table, widths)

    def add_list(self, token: dict[str, Any], level: int = 0) -> None:
        base = "ordered-list" if token.get("attrs", {}).get("ordered") else "unordered-list"
        style_name = f"{base}-{min(level + 1, 9)}"
        for item in token.get("children", []):
            for block in item.get("children", []):
                if block["type"] == "list":
                    self.add_list(block, level + 1)
                else:
                    paragraph = self.document.add_paragraph(style=style_name)
                    self.add_inline_nodes(paragraph, block.get("children", []))

    def add_code_block(self, token: dict[str, Any]) -> None:
        self.document.add_paragraph(
            token.get("raw", "").rstrip(), style="code-block"
        )

    def add_block_quote(self, token: dict[str, Any]) -> None:
        for child in token.get("children", []):
            if child["type"] in {"paragraph", "block_text"}:
                paragraph = self.document.add_paragraph(style="body")
                self.add_inline_nodes(paragraph, child.get("children", []))

    @staticmethod
    def _plain_text(nodes: Iterable[dict[str, Any]]) -> str:
        result = []
        for node in nodes:
            if "raw" in node and node["type"] in {"text", "codespan"}:
                result.append(node["raw"])
            result.append(DocxBuilder._plain_text(node.get("children", [])))
        return "".join(result)

    def consume(self, tokens: Iterable[dict[str, Any]]) -> None:
        handlers = {
            "heading": self.add_heading,
            "block_math": lambda token: self.add_block_math(token.get("raw", "")),
            "table": self.add_table,
            "list": self.add_list,
            "block_code": self.add_code_block,
            "block_quote": self.add_block_quote,
        }
        for token in tokens:
            if token["type"] == "paragraph":
                self.add_paragraph(token.get("children", []))
            elif token["type"] in handlers:
                handlers[token["type"]](token)
            elif token["type"] == "thematic_break":
                raise ValueError("Markdown thematic breaks are not rendered automatically")


def convert_markdown(
    input_path: str | Path,
    output_path: str | Path,
    config_path: str | Path,
) -> Path:
    input_path = Path(input_path)
    output_path = Path(output_path)
    metadata, markdown_source = parse_frontmatter(
        input_path.read_text(encoding="utf-8")
    )
    markdown = mistune.create_markdown(
        renderer="ast",
        plugins=["table", "math", "strikethrough"],
    )
    builder = DocxBuilder(load_config(config_path), input_path.parent)
    if "title" in metadata:
        builder.add_title(str(metadata["title"]))
    builder.consume(markdown(markdown_source))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    builder.document.save(output_path)
    return output_path
